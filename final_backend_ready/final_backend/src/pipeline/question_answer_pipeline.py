# src/pipeline/question_answer_pipeline.py
from src.utils.pydantic_compact import apply_pydantic_compat
apply_pydantic_compat()

import os
import io
import glob
import csv
import base64
import mimetypes
import requests
import traceback
import logging
from typing import List, Optional, Tuple

from langchain_community.vectorstores import FAISS

from src.components.rag_pipeline import RAGPipeline
from src.utils.synonym_expander import SynonymExpander

# ---------- Optional deps for PDF/image/table handling (auto-fallback if missing) ----------
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pdfplumber  # better table extraction from PDFs (optional)
except Exception:
    pdfplumber = None

try:
    import docx  # python-docx for .docx (optional)
except Exception:
    docx = None

# ---------- Directories (match main.py usage) ----------
UPLOAD_BASE = os.path.abspath(os.path.join(os.getcwd(), "uploaded_docs"))
VSTORE_BASE = os.path.abspath(os.path.join(os.getcwd(), "vectorstores"))

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
_PDF_EXT = ".pdf"
_DOCX_EXT = ".docx"
_TXT_EXTS = {".txt", ".md", ".rst"}
_CSV_EXTS = {".csv", ".tsv"}

# ---------- Small, safe chunking defaults ----------
_CHUNK_SIZE = 1000
_CHUNK_OVERLAP = 150
_MAX_TABLE_ROWS = 200  # don't explode the index with giant tables


def _safe_filename(name: str) -> str:
    """Strip any client path/backslashes and return just the filename."""
    return os.path.basename((name or "").replace("\\", "/"))


def _safe_strip_margin(s: str) -> str:
    """Remove leading spaces from a multi-line snippet for cleaner prompts."""
    return "\n".join(line.strip() for line in s.splitlines())


def _ext(fname: Optional[str]) -> str:
    return os.path.splitext(fname or "")[-1].lower()


def _find_uploaded_file(chat_id: str, stem: str) -> Optional[str]:
    """
    Locate the original uploaded file on disk by doc_id (stem).
    Looks under uploaded_docs/<chat_id>/<stem>.*
    """
    chat_dir = os.path.join(UPLOAD_BASE, chat_id)
    if not os.path.isdir(chat_dir):
        return None
    patt = os.path.join(chat_dir, f"{stem}.*")
    cand = sorted(glob.glob(patt))
    if cand:
        return cand[0]
    direct = os.path.join(chat_dir, stem)
    if os.path.exists(direct):
        return direct
    return None


def _pdf_text_length(pdf_path: str, max_pages: int = 5) -> int:
    """Return total length of extracted text from first N pages."""
    if not fitz:
        return 0
    try:
        doc = fitz.open(pdf_path)
        tlen = 0
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            tlen += len(page.get_text("text") or "")
        return tlen
    except Exception:
        return 0


def _pdf_to_images_b64(pdf_path: str, max_pages: int = 3, dpi: int = 220) -> List[str]:
    """Render first N pages of a PDF to PNG and return base64 list."""
    out: List[str] = []
    if not fitz:
        return out
    try:
        doc = fitz.open(pdf_path)
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            zoom = dpi / 72.0
            mtx = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mtx, alpha=False)
            png_bytes = pix.tobytes("png")
            out.append(base64.b64encode(png_bytes).decode("utf-8"))
    except Exception as e:
        print(f"⚠️ PDF render failed: {e}")
    return out


def _image_file_to_b64(path: str, target_min: int = 900) -> Optional[str]:
    """
    Read an image file and return JPEG/PNG base64.
    Optionally upscale smallest side for OCR readability.
    """
    try:
        if Image is None:
            with open(path, "rb") as f:
                raw = f.read()
            return base64.b64encode(raw).decode("utf-8")
        im = Image.open(path)
        im = im.convert("RGB")
        w, h = im.size
        if min(w, h) < target_min:
            scale = max(1.0, target_min / float(min(w, h)))
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=92, optimize=True)
        return base64.b64encode(buf.getvalue()).decode("utf-8")
    except Exception as e:
        print(f"⚠️ image read failed: {e}")
        return None


def _markdown_table(rows: List[List[str]]) -> str:
    """Convert a 2D list to a compact markdown table string."""
    if not rows:
        return ""
    headers = rows[0]
    headers = [str(h).strip() or f"col{i+1}" for i, h in enumerate(headers)]
    out = []
    out.append("| " + " | ".join(headers) + " |")
    out.append("| " + " | ".join("---" for _ in headers) + " |")
    for r in rows[1:_MAX_TABLE_ROWS]:
        out.append("| " + " | ".join(str(c).strip() for c in r) + " |")
    if len(rows) > _MAX_TABLE_ROWS:
        out.append(f"\n> Note: table truncated to first {_MAX_TABLE_ROWS} rows.\n")
    return "\n".join(out)


def _chunk_text(text: str, size: int = _CHUNK_SIZE, overlap: int = _CHUNK_OVERLAP) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        chunk = text[start:end]
        chunks.append(chunk)
        if end == len(text):
            break
        start = max(end - overlap, 0)
    return chunks


class QuestionAnswerPipeline:
    """
    High-level QA runner over your RAG pipeline.

    NEW:
    • Auto-train per document if vectorstore is missing:
      - Text docs: extract text + tables (PDF/DOCX/TXT/CSV)
      - Image/scanned docs: OCR via vision model (PDF pages → images, pure images)
    • Ensures answers even for table-heavy and image-heavy files.
    """
    def __init__(self):
        self.rag = RAGPipeline()

    # ------------------------------ Public helpers ------------------------------
    def sanitize_text(self, text: str) -> str:
        return text.encode("utf-8", "replace").decode("utf-8")

    # ------------------------------ Core run ------------------------------------
    def run(
        self,
        question: str,
        chat_id: str,
        document_id: Optional[str] = None,
        combine_docs: Optional[List[str]] = None,
        return_sources: bool = False,
    ):
        try:
            if not question:
                raise ValueError("No question provided.")
            if isinstance(question, list):
                question = " ".join(str(q) for q in question)
            elif not isinstance(question, str):
                raise TypeError("Invalid input: 'question' must be a string.")
            question = question.strip()
            if not question:
                raise ValueError("Question is empty after stripping.")

            expander = SynonymExpander()
            expanded_question = expander.find_similar_words(question)

            ql = question.lower()
            is_cost_question = any(
                kw in ql
                for kw in ["cost", "price", "charges", "amount", "budget", "fee", "bill", "expense"]
            )

            # ---------- Ensure training for combine-docs if needed ----------
            if combine_docs and len(combine_docs) > 0:
                print("🔗 Combine Mode: Ensuring indexes for selected documents...")
                ensured_docs = []
                for did in combine_docs:
                    s = os.path.splitext(_safe_filename(str(did)))[0]
                    upath = _find_uploaded_file(chat_id, s)
                    self._ensure_vectorstore(chat_id, s, upath)
                    ensured_docs.append(s)

                print("🔗 Combine Mode: Using selected documents only...")
                context, used_docs = self.rag.get_context(
                    question=expanded_question, chat_id=chat_id, document_id=None, combine_docs=ensured_docs, k=12
                )

                if not (context or "").strip():
                    raise RuntimeError("No relevant content found in the selected documents.")

                combined_contexts = self.sanitize_text(context)
                used_doc_lines = "\n".join(f"• {d}" for d in (used_docs or []))
                doc_hint = f"\nKnown document identifiers for citation:\n{used_doc_lines}\n" if used_docs else ""
                prompt = (
                    self.build_cost_compare_prompt(question, combined_contexts, doc_hint)
                    if is_cost_question
                    else self.build_multi_doc_compare_prompt(question, combined_contexts, doc_hint)
                )
                print(f"📄 Context from {len(used_docs)} selected documents being sent to model.")
                answer = self.call_ollama(prompt)
                return (answer, used_docs) if return_sources else answer

            # ---------- Single document path with smart training/router ----------
            if document_id:
                safe_key = _safe_filename(str(document_id))
                stem = os.path.splitext(safe_key)[0]
                uploaded_path = _find_uploaded_file(chat_id, stem)
                uploaded_ext = _ext(uploaded_path) if uploaded_path else None

                # Ensure vectorstore (this performs "training" if missing)
                self._ensure_vectorstore(chat_id, stem, uploaded_path)

                # Get context (works for both text & OCR’d content if we built it above)
                context, used_docs = self.rag.get_context(
                    question=expanded_question, chat_id=chat_id, document_id=stem, combine_docs=None, k=10
                )
                sanitized_context = self.sanitize_text(context or "")

                # Decide modality for answering:
                # 1) If it's a PDF with very low real text → use vision at answer-time (with tiny hint from RAG)
                # 2) If it's a pure image → use vision at answer-time
                # 3) Else → text/RAG
                go_vision = False
                images_b64: List[str] = []

                if uploaded_path:
                    if uploaded_ext == _PDF_EXT:
                        tlen = _pdf_text_length(uploaded_path, max_pages=5) if fitz else len(sanitized_context)
                        go_vision = (tlen < 200)
                        if go_vision:
                            images_b64 = _pdf_to_images_b64(uploaded_path, max_pages=3, dpi=220)
                    elif uploaded_ext in _IMAGE_EXTS:
                        go_vision = True
                        b = _image_file_to_b64(uploaded_path)
                        if b:
                            images_b64 = [b]
                    else:
                        if len(sanitized_context) < 150 and uploaded_ext in _CSV_EXTS | _TXT_EXTS | {_DOCX_EXT}:
                            go_vision = False  # rely on RAG for these formats
                else:
                    if len(sanitized_context) < 150:
                        go_vision = False  # no images we can attach

                if go_vision and images_b64:
                    prompt = self.build_vision_doc_prompt(
                        question=question,
                        doc_name=stem,
                        hint_text=sanitized_context[:1200]
                    )
                    print(f"🧭 Modality: VISION (images={len(images_b64)}) → {stem}")
                    answer = self.call_ollama(prompt, images_b64=images_b64)
                    if return_sources:
                        return answer, [stem]
                    return answer

                if not sanitized_context.strip():
                    raise RuntimeError("No relevant content found in the document.")

                prompt = self.build_single_doc_prompt(question, sanitized_context, stem)
                print(f"🧭 Modality: TEXT → {stem}")
                print(f"📄 Sending prompt to model from single document: {stem}")
                answer = self.call_ollama(prompt)

                if return_sources:
                    try:
                        vs_folder = os.path.join(VSTORE_BASE, chat_id, stem)
                        vectorstore = FAISS.load_local(
                            vs_folder, self.rag.embedding_model, allow_dangerous_deserialization=True
                        )
                        docs = vectorstore.similarity_search(expanded_question, k=5)
                        return answer, docs
                    except Exception:
                        return answer, [stem]
                return answer

            # ---------- All-docs path ----------
            print("ℹ️ No document selected; ensuring indexes for all chat documents and using all chat documents.")
            # Train any missing indexes under this chat
            chat_dir = os.path.join(UPLOAD_BASE, chat_id)
            if os.path.isdir(chat_dir):
                for path in sorted(glob.glob(os.path.join(chat_dir, "*"))):
                    stem = os.path.splitext(os.path.basename(path))[0]
                    self._ensure_vectorstore(chat_id, stem, path)

            context, used_docs = self.rag.get_context(
                question=expanded_question, chat_id=chat_id, document_id=None, combine_docs=[], k=12
            )
            if not context.strip():
                raise RuntimeError("No relevant content found across chat documents.")
            combined_contexts = self.sanitize_text(context)
            used_doc_lines = "\n".join(f"• {d}" for d in (used_docs or []))
            doc_hint = f"\nKnown document identifiers for citation:\n{used_doc_lines}\n" if used_docs else ""
            prompt = (
                self.build_cost_compare_prompt(question, combined_contexts, doc_hint)
                if is_cost_question
                else self.build_multi_doc_compare_prompt(question, combined_contexts, doc_hint)
            )
            answer = self.call_ollama(prompt)
            return (answer, used_docs) if return_sources else answer

        except Exception as e:
            print("❌ Pipeline Error: An unexpected error occurred. Full traceback below:")
            traceback.print_exc()
            error_type = type(e).__name__
            error_msg = str(e)
            full_error = f"❌ Backend Error: {error_type}('{error_msg}')"
            return (full_error, []) if return_sources else full_error

    # -------------------------------------------------------------------------
    # Training / Indexing helpers
    # -------------------------------------------------------------------------
    def _ensure_vectorstore(self, chat_id: str, stem: str, uploaded_path: Optional[str]) -> None:
        """
        If FAISS index for (chat_id, stem) is missing, build it from the file.
        - Text docs: extract text and tables.
        - Image/scanned: OCR via vision model to text+markdown tables.
        """
        vs_folder = os.path.join(VSTORE_BASE, chat_id, stem)
        index_file = os.path.join(vs_folder, "index.faiss")
        if os.path.exists(index_file):
            return  # already trained

        os.makedirs(vs_folder, exist_ok=True)

        if not uploaded_path or not os.path.exists(uploaded_path):
            print(f"⚠️ Cannot train: uploaded file for '{stem}' not found.")
            return

        ext = _ext(uploaded_path)
        print(f"🧩 Training index for '{stem}' ({ext}) ...")

        # Route: text-first extraction for known formats
        if ext in _TXT_EXTS | _CSV_EXTS | {_DOCX_EXT} | {_PDF_EXT}:
            # For PDF, decide if it's scanned vs. selectable
            if ext == _PDF_EXT and _pdf_text_length(uploaded_path, max_pages=5) < 200:
                # Scanned → OCR via vision
                text = self._ocr_pdf_to_text(uploaded_path)
            else:
                text = self._extract_text_and_tables(uploaded_path)
        elif ext in _IMAGE_EXTS:
            text = self._ocr_images_to_text([uploaded_path])
        else:
            # Unknown: best effort try text read; fall back to OCR if image-readable
            text = self._extract_text_and_tables(uploaded_path)
            if len(text.strip()) < 50:
                try:
                    if Image:
                        b64 = _image_file_to_b64(uploaded_path)
                        if b64:
                            text = self._ocr_b64s_to_text([b64])
                except Exception:
                    pass

        text = self.sanitize_text(text or "")
        if not text.strip():
            print(f"⚠️ No text extracted for '{stem}'. Skipping index build.")
            return

        chunks = _chunk_text(text)
        if not chunks:
            print(f"⚠️ No chunks produced for '{stem}'. Skipping index build.")
            return

        # Build FAISS using the same embedding model as RAGPipeline so downstream retrieval works
        embeddings = self.rag.embedding_model
        metadatas = [{"doc_id": stem, "source": stem}] * len(chunks)
        vectorstore = FAISS.from_texts(chunks, embedding=embeddings, metadatas=metadatas)
        vectorstore.save_local(vs_folder)
        print(f"✅ Trained and saved FAISS index at {vs_folder}")

    def _extract_text_and_tables(self, path: str) -> str:
        """
        Extracts text AND table content (as markdown tables) from various formats.
        """
        ext = _ext(path)
        try:
            if ext in _TXT_EXTS:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()

            if ext in _CSV_EXTS:
                rows: List[List[str]] = []
                with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
                    if ext == ".tsv":
                        reader = csv.reader(f, delimiter="\t")
                    else:
                        reader = csv.reader(f)
                    for i, row in enumerate(reader):
                        rows.append([str(c) for c in row])
                        if i > _MAX_TABLE_ROWS:
                            break
                return _markdown_table(rows)

            if ext == _DOCX_EXT and docx is not None:
                d = docx.Document(path)
                parts: List[str] = []
                for p in d.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)
                # tables
                for tbl in d.tables:
                    rows = []
                    for r in tbl.rows:
                        rows.append([cell.text.strip() for cell in r.cells])
                        if len(rows) > _MAX_TABLE_ROWS:
                            break
                    if rows:
                        parts.append(_markdown_table(rows))
                return "\n\n".join(parts)

            if ext == _PDF_EXT:
                parts: List[str] = []
                # 1) text via PyMuPDF
                if fitz is not None:
                    try:
                        doc = fitz.open(path)
                        for page in doc:
                            t = page.get_text("text") or ""
                            if t.strip():
                                parts.append(t.strip())
                    except Exception:
                        pass
                # 2) tables via pdfplumber (if available)
                if pdfplumber is not None:
                    try:
                        with pdfplumber.open(path) as pdf:
                            for p in pdf.pages:
                                tables = p.extract_tables() or []
                                for tbl in tables:
                                    if tbl:
                                        parts.append(_markdown_table(tbl))
                    except Exception:
                        pass
                return "\n\n".join(parts)

            # Fallback: try to read as text regardless of extension
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        except Exception as e:
            print(f"⚠️ Text extraction failed for '{path}': {e}")
            return ""

    def _ocr_pdf_to_text(self, pdf_path: str) -> str:
        """OCR a (likely scanned) PDF by rendering first pages to images and running vision-OCR."""
        b64s = _pdf_to_images_b64(pdf_path, max_pages=6, dpi=220)
        if not b64s:
            return ""
        return self._ocr_b64s_to_text(b64s)

    def _ocr_images_to_text(self, image_paths: List[str]) -> str:
        """OCR pure image files via vision model."""
        b64s: List[str] = []
        for p in image_paths:
            b = _image_file_to_b64(p)
            if b:
                b64s.append(b)
        if not b64s:
            return ""
        return self._ocr_b64s_to_text(b64s)

    def _ocr_b64s_to_text(self, images_b64: List[str]) -> str:
        """
        Runs a strict OCR prompt through the vision model to get plain text and markdown tables only.
        """
        prompt = (
            "You are an OCR engine. Read ALL text and tables from the attached page images.\n"
            "Return ONLY raw plaintext and markdown tables representing the content, no summaries, no commentary.\n"
            "If a table is present, reproduce it as a markdown table preserving cell values, units and currency symbols.\n"
            "Normalize nothing except whitespace; keep original spellings and numbers."
        )
        try:
            txt = self.call_ollama(prompt, images_b64=images_b64)
            return txt or ""
        except Exception as e:
            print(f"⚠️ Vision OCR failed: {e}")
            return ""

    # -------------------------------------------------------------------------
    # Prompt builders
    # -------------------------------------------------------------------------
    def _table_skills(self) -> str:
        return (
            "If the context includes tables (markdown tables, CSV-like lines, or PDF-extracted rows):\n"
            "• Treat each row as data; infer headers from surrounding text if missing.\n"
            "• Preserve units and numbers; normalize currencies but keep the original symbol (e.g., $/₹).\n"
            "• When quoting evidence, you may quote a cell or a short row snippet.\n"
            "• If page/section markers like 'p. 3' or 'Page: 3' appear near the text, include them in the table.\n"
        )

    def build_single_doc_prompt(self, question: str, context: str, doc_name: str) -> str:
        return f"""
You are a precise assistant answering ONLY from the provided document content (this may be selectable text or OCR from scanned/handwritten images).
Document Identifier: - {doc_name}
Document Content (verbatim snippets & table fragments):
-------------------------------------------------------
{context}
User Question:
--------------
{question}
Instructions:
-------------
- Use only the content above; do NOT invent facts.
- First infer the modality briefly in your mind: text OR image (scan/photo/handwritten). If image quality is poor, keep answers conservative.
- If the document appears to be an ID or a filled form, extract key fields when present:
• Aadhaar: name, dob_or_yob, gender, aadhaar_number, address
• PAN: name, father_name, dob (YYYY-MM-DD), pan_number (pattern AAAAA9999A)
• Ration Card: card_number, state, holder_name, family_members, address
• Generic form: form_title; applicant (name, mobile, email, address); ids_mentioned (aadhaar, pan, ration); fields (key→value); signatures_present (true/false); photo_present (true/false); dates (YYYY-MM-DD)
- Prefer short, structured bullet points or a tight paragraph.
- {_safe_strip_margin(self._table_skills())}
- Normalize dates to YYYY-MM-DD when possible; you may keep the raw date in parentheses if helpful.
- If any specific field is unreadable from the scan/handwriting, write "Unreadable" for that field.
- If the answer is not present, reply: "Not found in the provided document."
- Do not include file paths.
Output:
-------
Provide the best possible answer concisely. If helpful, include a tiny Markdown table.
""".strip()

    def build_vision_doc_prompt(self, question: str, doc_name: str, hint_text: str = "") -> str:
        hint_block = f"\nTiny OCR/Text hint (may be partial/noisy):\n---\n{hint_text}\n---\n" if hint_text else ""
        return f"""
You are an expert visual document analyst. You will receive 1–3 page images of a document.
Your job is to read all visible text, numbers, tables, stamps, signatures, logos, and charts/graphs.
Use ONLY what is visible in the images. If something is unclear or not present, say so explicitly.

Document: {doc_name}
{hint_block}
Question:
---------
{question}

Guidelines:
- Identify the document type if obvious (e.g., invoice, receipt, certificate, application, letter, notice, statement, form, chart, timetable, syllabus).
- If tables are present, read them carefully; preserve units and currency symbols exactly as shown (₹, Rs, $, USD, INR, %).
- For charts/graphs, read axis labels, legends, and the main trend; extract the key numeric points if readable.
- For forms, extract fields as key→value pairs. If a field is unreadable or missing, write "Unreadable" or "Not stated".
- If dates appear, normalize to YYYY-MM-DD where possible (keep the raw if ambiguous).
- Be concise and precise. Do NOT invent facts.

Output:
- A short, direct answer to the user's question.
- When helpful, include a tiny structured bullet list or a compact markdown table for extracted key-values or figures.
""".strip()

    def build_multi_doc_compare_prompt(self, question: str, combined_contexts: str, doc_hint: str = "") -> str:
        return f"""
You are comparing multiple documents (mix of text PDFs and scanned/handwritten images). Use ONLY the provided context (text snippets, OCR text, and model-attached images).
{doc_hint}
Context (verbatim snippets & table fragments from all selected docs):
--------------------------------------------------------------------
{combined_contexts}
User Question:
--------------
{question}
Output format (use ALL sections and keep headings exactly):
# Executive Summary
• 3–6 bullets with the main conclusions and most important differences across documents.
# Document Registry
List each doc with its inferred modality/type:
- Doc Name → modality: text|image, doc_type: aadhaar|pan|ration|form|invoice|receipt|certificate|letter|chart|other/unknown, notes (if any)
# Comparison Table
Create a compact markdown table with columns exactly:
| Criteria | Doc Name | Evidence (quote or paraphrase) | Page/Section |
- Include the most important 6–12 rows.
- Criteria can be Dates, Costs, KYC Fields (name/ids/address), Requirements, Clauses, Totals, Headings, etc.
- Doc Name must match the identifiers provided.
- Evidence must be short and specific.
# Gaps / Missing Info
• Bullets for unclear, missing, or contradictory parts (e.g., unreadable handwriting).
# Recommendation
• 1–2 sentences suggesting the best option or next steps, with trade-offs.
Rules:
- Do NOT hallucinate; if something is not found, say "Not stated".
- Normalize dates to YYYY-MM-DD where possible; keep raw in Evidence if useful.
- {_safe_strip_margin(self._table_skills())}
- Keep total under ~500 words.
""".strip()

    def build_cost_compare_prompt(self, question: str, combined_contexts: str, doc_hint: str = "") -> str:
        return f"""
You are a finance analyst extracting and comparing costs from multiple documents. Use ONLY the context below.
Documents may be text-based or OCR from scanned/handwritten images; if handwriting is unclear, be conservative.
{doc_hint}
Context (verbatim snippets & table fragments):
----------------------------------------------
{combined_contexts}
User Question:
--------------
{question}
Output format (use ALL sections and keep headings exactly):
# Executive Summary
• 3–5 bullets summarizing totals, major cost drivers, taxes/fees, notable differences, and any mismatches between computed and reported totals.
# Comparison Table
| Cost Item | Doc Name | Amount | Evidence (quote/paraphrase) | Page/Section |
- Amount must include the currency symbol exactly as shown (e.g., $, ₹, Rs, USD, INR) and any unit (per month, per hour).
- Show taxes and fees when explicitly stated (GST, CGST, SGST, IGST, TDS, surcharge, discount).
- If an item is presented as quantity × unit price, show the final line amount as written; if you compute a check, note it in "Notes & Assumptions".
- If a document lacks costs for an item, write "Not stated".
- If text is present but unreadable due to scan/handwriting, write "Unreadable".
# Notes & Assumptions
• Any normalization, ranges, or missing values.
• Briefly show any check you performed (for example, subtotal plus taxes equals reported total). If there is a mismatch, mention it here.
• Do not convert currencies using external rates; only normalize or convert if an explicit rate is provided in the context.
• Mention OCR or handwriting uncertainties if they affect amounts.
# Recommendation
• Brief guidance on the lowest TCO or most reasonable option with trade-offs, based only on the provided context.
Rules:
- Extract only explicit numeric costs exactly as written (examples: "$20,000", "₹50,000", "Rs 1,200", "USD 300").
- Do NOT guess numbers; if absent, say "Not stated". If illegible, say "Unreadable".
- Keep dates in YYYY-MM-DD if normalization is possible; keep raw text in Evidence if helpful.
- Do not include file paths.
- {_safe_strip_margin(self._table_skills())}
- Keep the total under ~500 words.
""".strip()

    # -------------------------------------------------------------------------
    # Robust Ollama caller with vision/chat support, retries, and fallback
    # -------------------------------------------------------------------------
    def _looks_vision_model(self, name: str) -> bool:
        n = (name or "").lower()
        return any(k in n for k in ["minicpm", "vl", "llava", "vision", "qwen2.5vl", "qwen-vl", "llama-vision"])

    def call_ollama(self, prompt: str, images_b64: Optional[List[str]] = None) -> str:
        """
        Robust Ollama caller:
        - Vision models (/ images provided) -> /api/chat
        - Text models -> /api/generate (fallbacks to /api/chat if needed)
        - Maps 'minicpm-hybrid' -> 'qwen2.5vl-gpu:latest'
        - Reads GPU/CTX from env; retries with smaller ctx and fallback model
        - Logs full Ollama error body
        """
        logger = logging.getLogger(__name__)

        # ---- Base URL normalization ----------------------------------------
        base = (getattr(self.rag, "ollama_url", None) or os.getenv("OLLAMA_BASE_URL") or "http://192.168.0.88:11434").rstrip("/")
        if base.endswith("/api/generate") or base.endswith("/api/chat"):
            base = base.rsplit("/api/", 1)[0]

        # ---- Model resolution ----------------------------------------------
        requested = (getattr(self.rag, "model_name", None) or os.getenv("OLLAMA_MODEL") or "qwen2.5vl-gpu:latest").strip()
        if "minicpm-hybrid" in requested.lower():
            requested = "qwen2.5-7b-sql:latest"
        fallback = (os.getenv("OLLAMA_MODEL_FALLBACK") or "llama3.1:8b").strip()

        # ---- Options / limits ----------------------------------------------
        try:
            num_ctx = int(os.getenv("OLLAMA_NUM_CTX", "4096"))
        except Exception:
            num_ctx = 4096
        try:
            num_gpu_layers = int(os.getenv("OLLAMA_NUM_GPU_LAYERS", "20"))
        except Exception:
            num_gpu_layers = 20

        system_text = (
            "You are a meticulous analyst. Use ONLY the provided context. "
            "If information is missing, explicitly say so. Keep answers precise."
        )

        safe_prompt = self.sanitize_text(prompt or "")
        if len(safe_prompt) > 12000:
            safe_prompt = safe_prompt[-12000:]

        def _opts(ctx: int) -> dict:
            return {"num_ctx": ctx, "num_gpu_layers": num_gpu_layers}

        def _post(path: str, payload: dict) -> dict:
            url = f"{base}{path}"
            print(f"📡 POST {url} model='{payload.get('model')}' stream={payload.get('stream', False)}")
            r = requests.post(url, json=payload, timeout=2000)
            if r.status_code >= 400:
                raise RuntimeError(f"Ollama {r.status_code}: {r.text}")
            return r.json()

        def _generate(mname: str, ctx: int) -> str:
            payload = {
                "model": mname,
                "prompt": safe_prompt,
                "stream": False,
                "keep_alive": "30m",
                "options": _opts(ctx),
            }
            j = _post("/api/generate", payload)
            resp = (j.get("response") or "").strip()
            if not resp and "error" in j:
                raise RuntimeError(f"Ollama error: {j['error']}")
            if not resp:
                raise RuntimeError("Received empty response from Ollama (/api/generate).")
            return resp

        def _chat(mname: str, ctx: int, imgs: Optional[List[str]]) -> str:
            messages = [{"role": "system", "content": system_text}]
            user_msg = {"role": "user", "content": safe_prompt}
            if imgs:
                user_msg["images"] = imgs
            messages.append(user_msg)
            payload = {
                "model": mname,
                "messages": messages,
                "stream": False,
                "options": _opts(ctx),
            }
            j = _post("/api/chat", payload)
            if isinstance(j.get("message"), dict):
                resp = (j["message"].get("content") or "").strip()
            else:
                resp = (j.get("response") or "").strip()
            if not resp and "error" in j:
                raise RuntimeError(f"Ollama error: {j['error']}")
            if not resp:
                raise RuntimeError("Received empty response from Ollama (/api/chat).")
            return resp

        def _is_vision(name: str) -> bool:
            return self._looks_vision_model(name) or bool(images_b64)

        primary = requested
        try:
            if _is_vision(primary):
                try:
                    return _chat(primary, num_ctx, images_b64)
                except Exception as e1:
                    logging.getLogger(__name__).warning("chat failed on %s: %s; retrying with smaller ctx", primary, e1)
                    return _chat(primary, max(1536, num_ctx // 2), images_b64)
            else:
                try:
                    return _generate(primary, num_ctx)
                except Exception as e1:
                    logging.getLogger(__name__).warning("generate failed on %s: %s; retrying via /api/chat", primary, e1)
                    return _chat(primary, max(1536, num_ctx // 2), images_b64)
        except Exception as e2:
            logging.getLogger(__name__).exception("Primary model '%s' failed; trying fallback '%s'", primary, fallback)
            try:
                return _generate(fallback, 2048)
            except Exception:
                return _chat(fallback, 1536, images_b64)
