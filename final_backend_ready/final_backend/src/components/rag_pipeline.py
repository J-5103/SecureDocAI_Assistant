# src/components/rag_pipeline.py
import os
import time
from typing import List, Tuple, Optional, Dict, Any
import io
import json
import base64
import re
import logging
from contextlib import nullcontext as _nullcontext

# --- Smart PDF/Image extraction ---
import fitz  # PyMuPDF
from PIL import Image, ImageOps

import pandas as pd
import docx  # python-docx
import requests


print("RAGPipeline module loaded from:", __file__)


# --- Deterministic bank name matching (for count/list queries) ---
_BANK_WORDS: Tuple[str, ...] = (
    "bank",
    "bank ltd",
    "co-operative bank",
    "cooperative bank",
    "sahakari bank",
    "nagrik sahakari bank",
)
# Optional torch (don't crash if it's not installed)
try:
    import torch  # type: ignore
    _TORCH_AVAILABLE = True
except Exception:
    torch = None  # type: ignore
    _TORCH_AVAILABLE = False

# Optional HuggingFace TrOCR (GPU preferred)
_TROCR_AVAILABLE = False
try:
    from transformers import TrOCRProcessor, VisionEncoderDecoderModel  # type: ignore
    _TROCR_AVAILABLE = True
except Exception:
    TrOCRProcessor = None  # type: ignore
    VisionEncoderDecoderModel = None  # type: ignore
    _TROCR_AVAILABLE = False


from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document

# Optional local reranker
try:
    from sentence_transformers import CrossEncoder  # type: ignore
    _RERANK_AVAILABLE = True
except Exception:
    CrossEncoder = None  # type: ignore
    _RERANK_AVAILABLE = False

from src.utils.synonym_expander import SynonymExpander

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

DEFAULT_K = 8
CHUNK_SIZE = 900
CHUNK_OVERLAP = 150

# ===================== Model & Pipeline Defaults =====================
# Text LLM (GPU) for text-based docs
TEXT_LLM_MODEL = os.environ.get("TEXT_LLM_MODEL", "qwen2.5vl-gpu:latest")
# Vision LLM (GPU) for vision/scanned docs (Ollama tag)
VLM_MODEL = os.environ.get("VLM_MODEL", "qwen2.5vl-gpu:latest")
# TrOCR model id (printed best for typed scans; for handwriting use '-handwritten')
TROCR_MODEL_ID = os.environ.get("TROCR_MODEL_ID", "microsoft/trocr-large-printed")

# Render/vision config
PDF_RENDER_DPI = int(os.environ.get("PDF_RENDER_DPI", "200"))
VLM_TILE_PX = int(os.environ.get("VLM_TILE_PX", "1024"))
VLM_TILE_OVERLAP = float(os.environ.get("VLM_TILE_OVERLAP", "0.15"))
VLM_MIN_TEXT_CHARS = int(os.environ.get("VLM_MIN_TEXT_CHARS", "80"))

# Control flags
USE_TROCR = os.environ.get("USE_TROCR", "1") not in ("0", "false", "False")
AUTO_VLM_ON_LOW_TEXT = os.environ.get("AUTO_VLM_ON_LOW_TEXT", "1") not in ("0", "false", "False")



class RAGPipeline:
    """
    RAG helper (updated):
      - OCR-less flow: text layer → VLM transcription (tiling) for scanned PDFs/images
      - FAISS + (optional) CrossEncoder rerank
      - Text LLM for grounded answers + Vision route for images/low-text PDFs
      - Strict, structured summary mode
    """
    
    def __init__(
        self,
        vector_store_path: str = "vectorstores",
        ollama_url: str = "http://192.168.0.88:11434",
        model_name: str = "qwen2.5vl-gpu:latest",                  # text LLM (GPU)
        vlm_model_name: str ="qwen2.5vl-gpu:latest",                   # vision model (GPU)
       # vision model
        request_timeout: int = 120, 
    
        # --- accuracy/perf toggles ---
        ocr_fallback: bool = False,              # OCR disabled (OCR-less path)
        vlm_transcribe_on_lowtext: bool = True,  # VLM to transcribe when text layer is short
        vlm_transcribe_min_chars: int = VLM_MIN_TEXT_CHARS,
        use_reranker: bool = True,
        reranker_model: str = "cross-encoder/ms-marco-MiniLM-L-6-v2",
        verify_pass: bool = False,
        # --- routing thresholds ---
        min_context_chars_for_text: int = 140,
        pdf_low_text_chars_probe: int = 100,
        auto_vlm_pages: int = 2,
    ):
        device = "cuda" if (_TORCH_AVAILABLE and torch.cuda.is_available()) else "cpu"  # type: ignore[attr-defined]
        print(f"Initializing RAGPipeline with device: {device}")

        # Embeddings
        self.embedding_model = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": device},
        )

        self.vector_store_path = vector_store_path
        self.ollama_url = ollama_url.rstrip("/")
        self.model_name = model_name
        self.vlm_model_name = vlm_model_name
        self.request_timeout = int(request_timeout)

        self.ocr_fallback = ocr_fallback
        self.vlm_transcribe_on_lowtext = vlm_transcribe_on_lowtext
        self.vlm_transcribe_min_chars = vlm_transcribe_min_chars
        self.use_reranker = use_reranker and _RERANK_AVAILABLE
        self.verify_pass = verify_pass

        self.min_context_chars_for_text = int(min_context_chars_for_text)
        self.pdf_low_text_chars_probe = int(pdf_low_text_chars_probe)
        self.auto_vlm_pages = int(auto_vlm_pages)

        self._reranker = None
        if self.use_reranker:
            try:
                self._reranker = CrossEncoder(reranker_model)
                print("✅ Reranker loaded:", reranker_model)
            except Exception as e:
                print("⚠️ Reranker load failed, falling back to vanilla similarity:", e)
                self.use_reranker = False

        # ---- TrOCR (lazy init) ----
        self._trocr_processor = None
        self._trocr_model = None
        self._trocr_device = device
        if USE_TROCR and _TROCR_AVAILABLE:
            try:
                self._init_trocr()
                print(f"✅ TrOCR ready: {TROCR_MODEL_ID} on {self._trocr_device}")
            except Exception as e:
                print("⚠️ TrOCR init failed:", e)
        else:
            print("ℹ️ TrOCR disabled (set USE_TROCR=1 and ensure transformers installed).")

        print("TEXT LLM:", self.model_name, " | VLM:", self.vlm_model_name)

    # -------------------- TrOCR helpers --------------------
    def _init_trocr(self):
        if self._trocr_processor and self._trocr_model:
            return
        try:
            # prefer fast image processor/tokenizer if available
            self._trocr_processor = TrOCRProcessor.from_pretrained(TROCR_MODEL_ID, use_fast=True)  # type: ignore
        except TypeError:
            # older transformers that don't accept `use_fast`
            self._trocr_processor = TrOCRProcessor.from_pretrained(TROCR_MODEL_ID)

        self._trocr_model = VisionEncoderDecoderModel.from_pretrained(TROCR_MODEL_ID)

        if _TORCH_AVAILABLE and torch.cuda.is_available():
            self._trocr_model = self._trocr_model.to("cuda")
            self._trocr_device = "cuda"
        else:
            self._trocr_device = "cpu"
        self._trocr_model.eval()

    def _trocr_ocr_image(self, pil_img: Image.Image, max_new_tokens: int = 256) -> str:
        if not (USE_TROCR and _TROCR_AVAILABLE and self._trocr_model and self._trocr_processor):
            return ""
        img = ImageOps.autocontrast(pil_img.convert("RGB"))
        pixel_values = self._trocr_processor(images=img, return_tensors="pt").pixel_values
        if _TORCH_AVAILABLE:
            pixel_values = pixel_values.to(self._trocr_device)  # type: ignore
        with (torch.no_grad() if _TORCH_AVAILABLE else _nullcontext()):
            generated_ids = self._trocr_model.generate(
                pixel_values,
                max_new_tokens=max_new_tokens,
                num_beams=2,
                early_stopping=True,
            )
        text = self._trocr_processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
        return text.strip()

    # -------------------- path utils --------------------

    def _chat_folder(self, chat_id: str) -> str:
        return os.path.join(self.vector_store_path, chat_id)

    def _doc_folder(self, chat_id: str, doc_name_or_id: str) -> str:
        base = os.path.splitext(doc_name_or_id)[0]
        return os.path.join(self._chat_folder(chat_id), base)

    def _load_vs(self, folder_path: str):
        index_file = os.path.join(folder_path, "index.faiss")
        if not os.path.exists(index_file):
            return None
        return FAISS.load_local(
            folder_path,
            self.embedding_model,
            allow_dangerous_deserialization=True,
        )
        # -------- helpers for resolving selected docs to vectorstore folders --------
    def _list_doc_folders(self, chat_id: str) -> Dict[str, str]:
        """Return {folder_basename: absolute_path} for chat's vectorstores."""
        base = self._chat_folder(chat_id)
        out: Dict[str, str] = {}
        if not os.path.isdir(base):
            return out
        for name in os.listdir(base):
            p = os.path.join(base, name)
            if os.path.isdir(p) and os.path.exists(os.path.join(p, "index.faiss")):
                out[name] = p
        return out

    # --- helper to normalize names coming from the UI / filesystem
    def _normalize_name(self, s: str) -> str:
        s = os.path.splitext(s)[0]
        s = s.strip().lower()
        s = re.sub(r"[\s\-_\.]+", " ", s)   # unify separators
        s = re.sub(r"\s+", " ", s)
        return s

    def _match_doc_folder(self, chat_id: str, requested: str) -> Optional[str]:
        """
        Robust mapping of a UI-selected name to an existing vectorstore folder.
        Handles case, spaces, hyphens/underscores, extensions, and fuzzy fallback.
        """
        from difflib import SequenceMatcher

        all_folders = self._list_doc_folders(chat_id)
        if not all_folders:
            return None

        reqn = self._normalize_name(requested)
        norm_map = {self._normalize_name(k): p for k, p in all_folders.items()}

        # exact normalized
        if reqn in norm_map:
            return norm_map[reqn]

        # begins-with / contains
        for nk, p in norm_map.items():
            if nk.startswith(reqn) or reqn.startswith(nk):
                return p
        for nk, p in norm_map.items():
            if reqn in nk or nk in reqn:
                return p

        # fuzzy fallback
        best_p, best_score = None, 0.0
        for nk, p in norm_map.items():
            score = SequenceMatcher(None, reqn, nk).ratio()
            if score > best_score:
                best_p, best_score = p, score
        if best_score >= 0.72:
            print(f"ℹ️ Fuzzy matched '{requested}' → '{best_p}' (score {best_score:.2f})")
            return best_p

        print(f"⚠️ Could not resolve selected doc: '{requested}' (normalized='{reqn}')")
        print(f"   Available docs: {list(all_folders.keys())}")
        return None


    # -------------------- small helpers --------------------

    def sanitize_text(self, text: str) -> str:
        return text.encode("utf-8", "replace").decode("utf-8")

    def detect_question_type(self, question: str) -> str:
        q = (question or "").lower()
        if any(k in q for k in ["cost", "price", "budget", "amount", "charges"]):
            return "cost"
        if any(k in q for k in ["summary", "summarize", "overview", "extract summary", "brief"]):
            return "summary"
        return "default"

    # ===================== Summary helpers (NEW) =====================

    # India-centric detectors
    _AADHAAR_RE = re.compile(r"\b(\d{4}\s?\d{4}\s?\d{4})\b")
    _MOBILE_RE  = re.compile(r"\b([6-9]\d{9})\b")
    _EMAIL_RE   = re.compile(r"\b([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})\b")
    _DATE_RE    = re.compile(
        r"\b((?:0?[1-9]|[12]\d|3[01])[-/\.](?:0?[1-9]|1[0-2]|[A-Za-z]{3,})[-/\.](?:19|20)\d{2})\b",
        re.IGNORECASE
    )
    _URL_RE     = re.compile(r"\b((?:https?://)?[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/\S*)?)\b")
    _LABEL_SEP_RE = re.compile(r"\s*[:\-–]\s*")

    def _kv_from_lines(self, text: str) -> Dict[str, str]:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        kv: Dict[str, str] = {}
        for ln in lines:
            if ":" in ln or " - " in ln or " – " in ln:
                parts = self._LABEL_SEP_RE.split(ln, maxsplit=1)
                if len(parts) == 2:
                    k, v = parts
                    k = re.sub(r"\s+", " ", k).strip().lower()
                    v = v.strip()
                    if k and v and k not in kv:
                        kv[k] = v
        return kv

    def _detect_entities_flags(self, text: str) -> Dict[str, Any]:
        emails  = list(dict.fromkeys(m.group(1) for m in self._EMAIL_RE.finditer(text)))
        phones  = list(dict.fromkeys(m.group(1) for m in self._MOBILE_RE.finditer(text)))
        aadhaar = list(dict.fromkeys(m.group(1).replace(" ", "") for m in self._AADHAAR_RE.finditer(text)))
        dates   = list(dict.fromkeys(m.group(1) for m in self._DATE_RE.finditer(text)))
        urls    = list(dict.fromkeys(m.group(1) for m in self._URL_RE.finditer(text)))
        kv      = self._kv_from_lines(text)

        tl = text.lower()
        flags = {
            "signatures_present": any(w in tl for w in ["signature", "signed by", "हस्ताक्षर", "સહી"]),
            "photo_present": any(w in tl for w in ["photo", "photograph", "passport size"]),
        }

        guess = {
            "name": kv.get("name") or kv.get("applicant") or kv.get("applicant name"),
            "address": kv.get("address") or kv.get("residential address"),
            "email": emails[0] if emails else None,
            "mobile": phones[0] if phones else None,
            "aadhaar": aadhaar[0] if aadhaar else None,
            "dates_found": dates,
            "urls_found": urls,
            "kv_pairs": kv,
            "flags": flags,
        }
        return guess

    def _get_src_meta(self, chat_id: str, doc_id: Optional[str]) -> Dict[str, Any]:
        """Collect file meta: name/type/pages/size_bytes if available."""
        src_path = self._load_source_meta(chat_id, doc_id)
        meta = {"name": None, "type": None, "pages": 0, "size_bytes": 0}
        if not src_path:
            return meta
        try:
            meta["name"] = os.path.basename(src_path)
            meta["size_bytes"] = os.path.getsize(src_path) if os.path.exists(src_path) else 0
            lower = src_path.lower()
            if lower.endswith(".pdf"):
                meta["type"] = "pdf"
                try:
                    d = fitz.open(src_path)
                    meta["pages"] = len(d)
                    d.close()
                except Exception:
                    pass
            elif lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
                meta["type"] = "image"
                meta["pages"] = 1
            else:
                meta["type"] = (os.path.splitext(lower)[1] or "").lstrip(".") or "file"
        except Exception:
            pass
        return meta

    def _build_summary_prompt(self, doc_text: str, file_meta: Dict[str, Any], hints: Dict[str, Any]) -> str:
        schema = (
            '{'
            '"file":{"name":"str","type":"pdf|image|file","pages":"int","size_bytes":"int"},'
            '"entities":{"name":"str|null","mobile":"str|null","email":"str|null","aadhaar":"str|null","address":"str|null",'
            '"dates_found":["str",...],"urls_found":["str",...]},'
            '"flags":{"signatures_present":"bool","photo_present":"bool"},'
            '"sections":["str",...],'
            '"issues":["str",...]'
            '}'
        )

        return (
            "You are a meticulous document analyst for Indian KYC/forms and technical PDFs.\n"
            "ALWAYS output TWO parts in EXACT order:\n"
            "1) A SINGLE-LINE JSON exactly matching this schema (NO newlines):\n"
            f"{schema}\n"
            "2) A section titled **Detailed Summary** in Markdown.\n\n"
            "HARD RULES:\n"
            "- Use ONLY the provided text. Do NOT hallucinate.\n"
            "- If something is not present, write 'Not found'.\n"
            "- Keep emails/phones/dates exactly as written.\n"
            "- Prefer the most legible/complete value when duplicates conflict.\n"
            "- Do not invent headings or values not grounded in text.\n\n"
            "WRITE THE **Detailed Summary** AS:\n"
            "## Document Overview\n"
            "- Purpose, scope, intended use (1–3 bullets)\n"
            "- Source file meta (name/type/pages/size if available)\n\n"
            "## Structure & Sections\n"
            "- List of visible sections/labels (bulleted)\n\n"
            "## Entities Extracted\n"
            "- Name / Mobile / Email / Aadhaar / Address / Other IDs & Dates\n\n"
            "## Key Facts by Section\n"
            "- For each section/label, list Field: Value as bullets\n\n"
            "## Tables/Lists (if any)\n"
            "- Convert to Markdown tables or nested bullets\n\n"
            "## Validation & Issues\n"
            "- Missing/invalid/inconsistent fields\n\n"
            f"FILE META: {json.dumps(file_meta, ensure_ascii=False)}\n"
            f"DETECTED HINTS: {json.dumps(hints, ensure_ascii=False)}\n\n"
            "DOCUMENT TEXT (verbatim extract; do NOT summarize here):\n"
            f"\"\"\"{doc_text[:12000]}\"\"\"\n"
        )

    def _parse_json_and_markdown(self, raw: str) -> Tuple[str, str]:
        if not raw:
            return "", ""
        parts = raw.strip().splitlines()
        for i, ln in enumerate(parts):
            s = ln.strip()
            if s.startswith("{") and s.endswith("}"):
                return s, "\n".join(parts[i+1:]).strip()
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            json_line = re.sub(r"\s+", " ", m.group(0)).strip()
            md = (raw[:m.start()] + raw[m.end():]).strip()
            return json_line, md
        return "", raw.strip()

    # -------------------- Vision helpers (OCR-less) --------------------

    def _pil_to_b64(self, img: Image.Image) -> str:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return base64.b64encode(buf.getvalue()).decode("utf-8")

    def _tile_image(self, img: Image.Image) -> List[Image.Image]:
        w, h = img.size
        tile = max(256, VLM_TILE_PX)
        overlap = int(tile * VLM_TILE_OVERLAP)
        step = max(1, tile - overlap)
        tiles: List[Image.Image] = []
        for y in range(0, h, step):
            for x in range(0, w, step):
                box = (x, y, min(x + tile, w), min(y + tile, h))
                crop = img.crop(box)
                if crop.width >= 128 and crop.height >= 128:
                    tiles.append(crop)
            if y + tile >= h:
                break
        return tiles or [img]

    def _vlm_generate(self, prompt: str, images_b64: List[str], timeout: Optional[int] = None) -> str:
        """Ollama generate → chat fallback."""
        timeout = timeout or self.request_timeout
        payload = {
            "model": self.vlm_model_name,
            "prompt": prompt,
            "images": images_b64,
            "stream": False,
        }
        try:
            r = requests.post(f"{self.ollama_url}/api/generate", json=payload, timeout=timeout)
            r.raise_for_status()
            data = r.json()
            out = (data.get("response") or "").strip()
            if out:
                return out
        except Exception:
            pass
        # chat fallback
        try:
            chat_payload = {
                "model": self.vlm_model_name,
                "messages": [{"role": "user", "content": prompt, "images": [f"data:image/png;base64,{b}" for b in images_b64]}],
                "stream": False,
            }
            r2 = requests.post(f"{self.ollama_url}/api/chat", json=chat_payload, timeout=timeout)
            r2.raise_for_status()
            data2 = r2.json()
            return (data2.get("message", {}).get("content") or "").strip()
        except Exception:
            return ""

    def _vlm_transcribe_two_pass(self, pil_img: Image.Image) -> str:
        """Tile → two-pass prompts → merge (raw lines + CSV tables)."""
        img2 = ImageOps.autocontrast(pil_img.convert("RGB"))
        tiles = self._tile_image(img2)
        b64 = [self._pil_to_b64(t) for t in tiles]

        prompt_a = (
            "Transcribe ALL readable text lines from these image tiles. "
            "Output plain text only, one line per bullet, preserve numbers, case and punctuation."
        )
        text_a = self._vlm_generate(prompt_a, b64) or ""

        prompt_b = (
            "If any tables are visible in these image tiles, output ONLY their CSV (no commentary). "
            "If none, return an empty string."
        )
        text_b = self._vlm_generate(prompt_b, b64) or ""

        merged = (text_a.strip() + "\n" + text_b.strip()).strip()
        return self.sanitize_text(re.sub(r"\s+\n", "\n", merged))

    def _pdf_page_as_pil(self, page: fitz.Page, dpi: int = PDF_RENDER_DPI) -> Image.Image:
        mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    def ask_vlm(self, prompt: str, image_paths: List[str]) -> str:
        """Compatibility wrapper (kept for your existing calls)."""
        imgs_b64 = []
        for p in image_paths:
            with open(p, "rb") as f:
                imgs_b64.append(base64.b64encode(f.read()).decode("utf-8"))
        return self._vlm_generate(prompt, imgs_b64)

    def vlm_transcribe_page(self, page_png_path: str) -> str:
        with Image.open(page_png_path) as im:
            return self._vlm_transcribe_two_pass(im)

    # -------------------- extraction --------------------

    def extract_text_from_doc(self, file_path: str) -> List[Document]:
        """Extract text from various document types with OCR-less hybrid strategy."""
        filename = os.path.basename(file_path).replace(os.path.splitext(file_path)[1], "")
        documents: List[Document] = []
        try:
            lower = file_path.lower()

            if lower.endswith(".pdf"):
                print(f"Starting Smart PDF text extraction from: {filename}")
                doc = fitz.open(file_path)
                for i, page in enumerate(doc):
                    page_text = (page.get_text() or "").strip()
                    # If PDF has a usable text layer -> use it (TEXT path)
                    if len(page_text) >= VLM_MIN_TEXT_CHARS:
                        safe_text = self.sanitize_text(page_text)
                        documents.append(
                            Document(page_content=safe_text, metadata={"source": filename, "page": i + 1})
                        )
                    else:
                        # Vision/scanned page -> render + TrOCR (VISION path)
                        pil_img = self._pdf_page_as_pil(page, dpi=PDF_RENDER_DPI)
                        ocr_text = self._trocr_ocr_image(pil_img) if USE_TROCR else ""
                        if ocr_text.strip():
                            safe_text = self.sanitize_text(ocr_text)
                            documents.append(
                                Document(page_content=safe_text, metadata={"source": filename, "page": i + 1})
                            )
                doc.close()

            elif lower.endswith((".docx", ".doc")):
                d = docx.Document(file_path)
                parts: List[str] = []
                for para in d.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
                for t in d.tables:
                    for r in t.rows:
                        row = [cell.text.strip() for cell in r.cells]
                        if any(row):
                            parts.append(" | ".join(row))
                if parts:
                    safe_text = self.sanitize_text("\n".join(parts))
                    documents.append(Document(page_content=safe_text, metadata={"source": filename}))

            elif lower.endswith((".xlsx", ".xls", ".csv")):
                if lower.endswith((".xlsx", ".xls")):
                    df = pd.read_excel(file_path)
                else:
                    df = pd.read_csv(file_path)
                text = df.to_csv(index=False)
                if text.strip():
                    safe_text = self.sanitize_text(text)
                    documents.append(Document(page_content=safe_text, metadata={"source": filename}))

            elif lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
                # Vision images -> TrOCR (VISION path)
                im = Image.open(file_path).convert("RGB")
                ocr_text = self._trocr_ocr_image(im) if USE_TROCR else ""
                if ocr_text.strip():
                    safe_text = self.sanitize_text(ocr_text)
                    documents.append(Document(page_content=safe_text, metadata={"source": filename}))

            else:
                raise ValueError(f"Unsupported file type: {file_path}")

# (rest unchanged)

        except Exception as e:
            print(f"❌ Error reading {filename}: {e}")
            raise

        print(f"Text extraction complete for {filename}. Total documents: {len(documents)}")
        return documents

    def _infer_chat_id_from_pdf_path(self, file_path: Optional[str]) -> Optional[str]:
        if not file_path:
            return None
        parts = os.path.normpath(file_path).split(os.sep)
        try:
            up_idx = parts.index("uploaded_docs")
            return parts[up_idx + 1] if up_idx + 1 < len(parts) else None
        except ValueError:
            return None

    # -------------------- indexing --------------------

    def create_vectorstore(
        self,
        file_path: Optional[str] = None,
        combined_text: Optional[str] = None,
        vector_store_path: Optional[str] = None,
        metadata: Optional[dict] = None,
    ):
        filename = os.path.basename(file_path or "manual_input").replace(
            os.path.splitext(file_path or "")[1], ""
        )
        try:
            if not combined_text and not file_path:
                raise ValueError("Either 'file_path' or 'combined_text' must be provided.")

            inferred_chat = self._infer_chat_id_from_pdf_path(file_path)
            default_folder = (
                os.path.join(self.vector_store_path, inferred_chat, filename)
                if inferred_chat
                else os.path.join(self.vector_store_path, filename)
            )
            vectorstore_folder = vector_store_path or default_folder
            os.makedirs(vectorstore_folder, exist_ok=True)

            print(f"Starting vectorstore creation for: {filename} at {vectorstore_folder}")
            start_time = time.time()

            if combined_text:
                safe_text = self.sanitize_text(combined_text)
                docs = [Document(page_content=safe_text, metadata={"source": filename})]
            else:
                docs = self.extract_text_from_doc(file_path)  # type: ignore[arg-type]

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
            )
            chunks = splitter.split_documents(docs)
            print(f"Text splitting complete. Total chunks: {len(chunks)}")

            # Batch embedding
            batch_size = 100
            all_embeddings = []
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i : i + batch_size]
                embeddings = self.embedding_model.embed_documents([c.page_content for c in batch])
                all_embeddings.extend(embeddings)
                time.sleep(0.01)

            texts = [c.page_content for c in chunks]
            text_embeddings = list(zip(texts, all_embeddings))
            metadatas = [{"source": filename, "page": c.metadata.get("page", 1)} for c in chunks]

            vectorstore = FAISS.from_embeddings(
                text_embeddings=text_embeddings,
                embedding=self.embedding_model,
                metadatas=metadatas,
            )
            vectorstore.save_local(vectorstore_folder)

            # save source info for later
            try:
                with open(os.path.join(vectorstore_folder, "_source.json"), "w", encoding="utf-8") as f:
                    json.dump({"file_path": file_path, "filename": filename}, f)
            except Exception:
                pass

            elapsed = time.time() - start_time
            print(f"✅ Vectorstore created at: {vectorstore_folder} in {elapsed:.2f} seconds")

            return {
                "status": "ready",
                "document_id": os.path.splitext(filename)[0],
                "vectorstore_path": vectorstore_folder,
            }
        except Exception as e:
            print(f"❌ create_vectorstore failed for {filename}: {e}")
            raise

    def index_document(self, file_path: str, namespace: str, out_dir: str) -> Dict[str, Any]:
        return self.create_vectorstore(file_path=file_path, vector_store_path=out_dir)

    def search(self, query: str, namespace: str, index_dir: str, k: int = DEFAULT_K) -> List[Dict[str, Any]]:
        vs = self._load_vs(index_dir)
        if not vs:
            return []
        docs_scores = vs.similarity_search_with_score(query, k=k)
        return [
            {"text": d.page_content, "score": float(s), "source": d.metadata.get("source"), "page": d.metadata.get("page")}
            for d, s in docs_scores
        ]

    # -------------------- retrieval + rerank --------------------

    # NEW: LLM-based expansion to capture paraphrases and related concepts
    def _expand_query_llm(self, question: str) -> str:
        """
        Uses the text LLM to extract key terms/synonyms/related concepts so retrieval
        doesn't depend on exact wording. Returns a space-joined string of phrases.
        """
        prompt = (
            "Extract the core intent, keywords, synonyms and related concepts for the question. "
            "Return comma-separated short phrases only; no sentences.\n\n"
            f"Question: {question}\nPhrases:"
        )
        try:
            r = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": self.model_name,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.1, "num_predict": 128},
                },
                timeout=min(self.request_timeout, 60),
            )
            r.raise_for_status()
            raw = (r.json().get("response") or "").strip()
            terms = re.split(r"[,\n]+", raw)
            terms = [t.strip() for t in terms if t.strip()]
            return " ".join(terms)
        except Exception:
            return ""

    def _expanded_query(self, question: str) -> str:
        intent = self._parse_intent(question)
        expander = SynonymExpander()
        syn = expander.find_similar_words(question)
        llm_terms = self._expand_query_llm(question)
        add: List[str] = []
        if intent.get("startswith"):
            add.append(f"starts with {intent['startswith']}")
        if intent.get("location"):
            add += [intent["location"], f"{intent['location']} district", "list of banks"]
        combo = f"{question} {syn} {llm_terms} {' '.join(add)}".strip()
        return re.sub(r"\s+", " ", combo)


    def _topk_from_vs(self, vs, query: str, k: int) -> List[Tuple[Document, float]]:
        return vs.similarity_search_with_score(query, k=max(1, k))

    def _rerank_candidates(self, question: str, pairs: List[Tuple[Document, float]], final_k: int) -> List[Document]:
        if not self.use_reranker or not self._reranker:
            pairs.sort(key=lambda x: x[1])  # FAISS distance asc
            return [d for d, _ in pairs[:final_k]]

        texts = [(question, d.page_content) for d, _ in pairs]
        scores = self._reranker.predict(texts)  # higher better
        ranked = sorted(zip(scores, [d for d, _ in pairs]), key=lambda t: t[0], reverse=True)
        return [d for _, d in ranked[:final_k]]
    
    # ---------- INTENT + DETERMINISTIC FILTERS (NEW) ----------

    
    def _parse_intent(self, q: str) -> Dict[str, Any]:
        ql = (q or "").lower()
        intent = {
            "wants_count": any(w in ql for w in ["how many", "kitne", "count", "no. of", "number of"]),
            "startswith": None,
            "location": None,
        }
        # "which is start from M" / "start with P"
        m = re.search(r"(start\s+(?:with|from)\s+)([a-z])", ql)
        if m:
            intent["startswith"] = m.group(2).upper()

        # "in Ahmedabad District" / "at <place>" / "in <city>"
        m2 = re.search(r"\b(?:in|at)\s+([a-z][a-z\s\-]+?)(?:\s+district|\s+city|\s+state|$)", ql)
        if m2:
            intent["location"] = m2.group(1).strip().title()

        return intent

    def _candidate_lines_from_context(self, ctx: str) -> List[str]:
        # Context aise aata hai: "1. ....\n2. ...." -> lines nikalo
        raw = [re.sub(r"^\s*\d+\.\s*", "", ln).strip() for ln in ctx.splitlines()]
        lines = []
        for ln in raw:
            if not ln:
                continue
            # Chunks ke andar multiple lines ho sakti; split further
            for seg in re.split(r"[•\-\u2022]|\\n", ln):
                s = seg.strip()
                if s:
                    lines.append(s)
        return lines

    def _filter_banks(self, lines: List[str], startswith: Optional[str], location: Optional[str]) -> List[str]:
        out = []
        for ln in lines:
            lnl = ln.lower()
            if not any(w in lnl for w in _BANK_WORDS):
                continue
            if startswith and not re.match(rf"^{re.escape(startswith)}", ln.strip(), flags=re.IGNORECASE):
                continue
            if location and (location.lower() not in lnl):
                continue
            # cleanup tail punctuation
            clean = re.sub(r"\s*[,:;.\-–—]+$", "", ln).strip()
            out.append(clean)
        # unique, preserve order
        seen = set(); uniq = []
        for x in out:
            k = x.casefold()
            if k in seen: continue
            seen.add(k); uniq.append(x)
        return uniq


    def get_context_from_single_doc(
        self,
        question: str,
        chat_id: str,
        document_id: str,
        top_k: int = DEFAULT_K,
    ) -> Tuple[str, List[str], str]:
        question_type = self.detect_question_type(question)
        expanded_query = self._expanded_query(question)

        folder_path = self._doc_folder(chat_id, document_id)
        vs = self._load_vs(folder_path)
        if not vs:
            print(f"❌ Vectorstore not found for: {folder_path}")
            return "", [], question_type

        initial_k = max(20, top_k * 3)
        pairs = self._topk_from_vs(vs, expanded_query, initial_k)
        if not pairs:
            return "", [], question_type

        best_docs = self._rerank_candidates(question, pairs, final_k=top_k)
        numbered = []
        for idx, doc in enumerate(best_docs, start=1):
            doc_text = self.sanitize_text(doc.page_content.strip())
            numbered.append(f"{idx}. {doc_text}")

        ctx = "\n".join(numbered)
        return ctx, [os.path.basename(folder_path)], question_type

    def get_context_from_multiple_docs(
        self,
        question: str,
        chat_id: str,
        document_names: Optional[List[str]] = None,
        top_k: int = DEFAULT_K,
    ) -> Tuple[str, List[str], str]:
        print("🔍 Loading vectorstores across multiple docs...")
        question_type = self.detect_question_type(question)

        base_folder = self._chat_folder(chat_id)
        if not os.path.exists(base_folder):
            print(f"❌ Chat folder not found: {base_folder}")
            return "", [], question_type

        expanded_query = self._expanded_query(question)

        # --- resolve selected folders robustly ---
        # --- resolve selected folders robustly (use ALL that resolve) ---
        if document_names:
            target_folders: List[str] = []
            unresolved: List[str] = []
            for name in document_names:
                match = self._match_doc_folder(chat_id, name)
                if match:
                    target_folders.append(match)
                else:
                    unresolved.append(name)

            print(f"🧩 UI selected: {document_names}")
            print(f"✅ Resolved to folders: {target_folders}")
            if unresolved:
                print(f"⚠️ Unresolved selections (no vectorstore found): {unresolved}")

            if not target_folders:
                return "", [], question_type
        else:
            target_folders = list(self._list_doc_folders(chat_id).values())
            if not target_folders:
                return "", [], question_type


        # --- retrieve per-doc, then rerank globally ---
        global_pairs: List[Tuple[Document, float, str]] = []
        per_doc_k = max(8, top_k)  # fair share from each doc

        for folder_path in target_folders:
            vs = self._load_vs(folder_path)
            if not vs:
                print(f"⚠️ Skipping missing/empty VS: {folder_path}")
                continue
            try:
                pairs = self._topk_from_vs(vs, expanded_query, per_doc_k)
                for d, s in pairs:
                    global_pairs.append((d, float(s), folder_path))
            except Exception as e:
                print(f"❌ Failed search in {folder_path}: {e}")

        if not global_pairs:
            return "", [], question_type

        # ensure at least a couple chunks per selected doc
        by_doc: Dict[str, List[Tuple[Document, float]]] = {}
        for d, s, fp in global_pairs:
            by_doc.setdefault(fp, []).append((d, s))

        seeds: List[Document] = []
        min_each = max(2, top_k // max(1, len(by_doc)))
        for fp, pairs in by_doc.items():
            pairs.sort(key=lambda x: x[1])  # FAISS distance asc
            seeds.extend([d for d, _ in pairs[:min_each]])

        # fill the rest with top global matches
        remaining = [d for (d, _s, _fp) in global_pairs if d not in seeds][: max(1, top_k * 3 - len(seeds))]
        candidates = seeds + remaining

        # rerank (or distance sort fallback handled inside)
        best_docs = self._rerank_candidates(question, [(d, 0.0) for d in candidates], final_k=top_k)

        numbered: List[str] = []
        used_docs: List[str] = []
        for idx, d in enumerate(best_docs, start=1):
            doc_text = self.sanitize_text(d.page_content.strip())
            numbered.append(f"{idx}. {doc_text}")
            used_docs.append(d.metadata.get("source") or "doc")

        # dedupe used_docs, preserve order
        seen = set()
        used_docs = [x for x in used_docs if not (x in seen or seen.add(x))]

        ctx = "\n".join(numbered)
        print(f"🧩 Combined from docs: {used_docs}")
        return ctx, used_docs, question_type


    def get_context(
        self,
        question: str,
        chat_id: str,
        document_id: Optional[str] = None,
        combine_docs: Optional[List[str]] = None,
        k: int = DEFAULT_K,
    ) -> Tuple[str, List[str]]:
        if document_id:
            ctx, used, _ = self.get_context_from_single_doc(question, chat_id, document_id, top_k=k)
            return ctx, used
        ctx, used, _ = self.get_context_from_multiple_docs(question, chat_id, document_names=combine_docs, top_k=k)
        return ctx, used

    # -------------------- summary helpers --------------------

    def _load_source_meta(self, chat_id: str, doc_id: Optional[str]) -> Optional[str]:
        if not doc_id:
            return None
        folder = self._doc_folder(chat_id, doc_id)
        meta = os.path.join(folder, "_source.json")
        if os.path.exists(meta):
            try:
                with open(meta, "r", encoding="utf-8") as f:
                    return json.load(f).get("file_path")
            except Exception:
                return None
        return None

    def _probe_pdf_text_chars(self, pdf_path: str, probe_pages: int = 3) -> int:
        try:
            doc = fitz.open(pdf_path)
            n = min(probe_pages, len(doc))
            total = 0
            for i in range(n):
                total += len((doc[i].get_text() or "").strip())
            doc.close()
            return total
        except Exception:
            return 0

    def _render_pdf_to_images(self, chat_id: str, doc_id: Optional[str], max_pages: int = 4, dpi: int = 260) -> List[str]:
        """Render first N pages of the original PDF to PNGs for vision."""
        src = self._load_source_meta(chat_id, doc_id)
        if not src or not str(src).lower().endswith(".pdf"):
            return []
        try:
            pdf = fitz.open(src)
            paths = []
            for i in range(min(max_pages, len(pdf))):
                pix = pdf[i].get_pixmap(dpi=dpi)
                outp = f"{src}.p{i+1:03d}.png"
                pix.save(outp)
                paths.append(outp)
            pdf.close()
            return paths
        except Exception as e:
            print("render_pdf_to_images failed:", e)
            return []

    def _summary_prompt_vlm(self, question: str) -> str:
        return (
            "You are a meticulous document analyst. Read the attached form/pages and produce a CLEAN SUMMARY.\n"
            "Rules:\n"
            "• Use only what is clearly visible in the images.\n"
            "• If a field is unreadable, write: Not clearly visible.\n"
            "• Return concise bullet points; keep emails/phones exactly as printed.\n\n"
            f"Question: {question}\n"
            "Return markdown only."
        )

    def summarize_document(self, question: str, chat_id: str,
                           document_id: Optional[str], combine_docs: Optional[List[str]]) -> str:
        target_doc = document_id or (combine_docs[0] if combine_docs else None)
        file_meta = self._get_src_meta(chat_id, target_doc)

        images = self._render_pdf_to_images(chat_id, target_doc, max_pages=self.auto_vlm_pages, dpi=280)

        doc_text_for_hints = ""
        if images:
            try:
                vlm_plain = self.ask_vlm(
                    "Transcribe all readable text from the page images. Return PLAIN TEXT only, no bullets, no headings.",
                    images
                )
                doc_text_for_hints = (vlm_plain or "").strip()
            except Exception:
                doc_text_for_hints = ""
        if not doc_text_for_hints:
            context, _used = self.get_context(
                question=question, chat_id=chat_id,
                document_id=document_id, combine_docs=combine_docs, k=DEFAULT_K
            )
            doc_text_for_hints = (context or "").strip()

        hints = self._detect_entities_flags(doc_text_for_hints)

        if images:
            prompt = self._build_summary_prompt(doc_text_for_hints, file_meta, hints)
            print(f"🖼️ Summary via VLM on {len(images)} page image(s)")
            raw = (self.ask_vlm(prompt, images) or "").strip()
            json_line, md = self._parse_json_and_markdown(raw)
        else:
            context, used_docs = self.get_context(
                question=question, chat_id=chat_id,
                document_id=document_id, combine_docs=combine_docs, k=DEFAULT_K
            )
            if not context:
                minimal = {
                    "file": file_meta,
                    "entities": {"name": None, "mobile": None, "email": None, "aadhaar": None,
                                 "address": None, "dates_found": [], "urls_found": []},
                    "flags": {"signatures_present": hints.get("flags", {}).get("signatures_present", False),
                              "photo_present": hints.get("flags", {}).get("photo_present", False)},
                    "sections": [],
                    "issues": ["No extractable text found in document/context"]
                }
                return json.dumps(minimal, ensure_ascii=False) + "\n\n**Detailed Summary**\n\nNot found in the provided document."

            text_for_prompt = context
            prompt = self._build_summary_prompt(text_for_prompt, file_meta, hints)
            payload = {
                "model": self.model_name,
                "prompt": prompt,
                "stream": False,
                "keep_alive": "10m",
                "options": {"temperature": 0.1, "num_predict": 700, "num_ctx": 4096},
            }
            print(f"➡️  TEXT LLM (summary): {self.model_name}")
            r = requests.post(f"{self.ollama_url}/api/generate", json=payload, timeout=self.request_timeout)
            r.raise_for_status()
            raw = (r.json().get("response") or "").strip()
            json_line, md = self._parse_json_and_markdown(raw)

        if not json_line:
            fallback = {
                "file": file_meta,
                "entities": {
                    "name": hints.get("name"),
                    "mobile": hints.get("mobile"),
                    "email": hints.get("email"),
                    "aadhaar": hints.get("aadhaar"),
                    "address": hints.get("address"),
                    "dates_found": hints.get("dates_found", []),
                    "urls_found": hints.get("urls_found", [])
                },
                "flags": {
                    "signatures_present": hints.get("flags", {}).get("signatures_present", False),
                    "photo_present": hints.get("flags", {}).get("photo_present", False),
                },
                "sections": sorted(list(hints.get("kv_pairs", {}).keys()))[:20],
                "issues": ["LLM output parsing failed; returned best-effort detectors"]
            }
            json_line = json.dumps(fallback, ensure_ascii=False)
        if not md:
            md = "**Detailed Summary**\n\n(LLM did not return a Markdown section.)"

        return f"{json_line}\n\n{md}"

    # -------------------- answer --------------------

    def _verify_answer(self, question: str, context: str, answer: str) -> str:
        prompt = (
            "You are verifying an answer strictly against the given context. "
            "If the answer is fully supported by the context, return the same answer. "
            "If unsupported or contradictory, reply exactly: \"I don't know\".\n\n"
            f"CONTEXT:\n{context}\n\nQUESTION: {question}\nANSWER TO VERIFY: {answer}\nFINAL:"
        )
        payload = {
            "model": self.model_name,
            "prompt": prompt,
            "stream": False,
            "keep_alive": "10m",
            "options": {"temperature": 0.0, "num_predict": 200, "num_ctx": 3072},
        }
        try:
            print(f"➡️  TEXT LLM (verify): {self.model_name}")
            r = requests.post(f"{self.ollama_url}/api/generate", json=payload, timeout=self.request_timeout)
            r.raise_for_status()
            out = (r.json().get("response") or "").strip()
            return out or answer
        except Exception:
            return answer

    def _auto_vlm_from_pdf(self, question: str, chat_id: str, doc_id: Optional[str]) -> Optional[str]:
        src = self._load_source_meta(chat_id, doc_id)
        if not src or not str(src).lower().endswith(".pdf"):
            return None
        try:
            pdf = fitz.open(src)
            max_pages = min(self.auto_vlm_pages, len(pdf))
            img_paths = []
            for i in range(max_pages):
                pix = pdf[i].get_pixmap(dpi=220)
                outp = f"{src}.page_{i+1:03d}.png"
                pix.save(outp)
                img_paths.append(outp)
            pdf.close()
            print(f"🖼️ Auto VLM on PDF pages: {img_paths}")
            vlm_prompt = (
                "Answer strictly from these document pages. "
                "Interpret the user's intent semantically — do not require exact wording matches. "
                "If not visible, reply exactly: \"I don't know\".\n\n"
                f"Question: {question}"
            )
            return (self.ask_vlm(vlm_prompt, img_paths) or "").strip()
        except Exception as e:
            print("Auto VLM failed:", e)
            return None

    def answer_question(self, question: str, chat_id: str,
                    document_id: Optional[str] = None,
                    combine_docs: Optional[List[str]] = None,
                    images: Optional[List[str]] = None) -> str:
        try:
            qtype = self.detect_question_type(question)
            note = ("\n\n*Note:* I derive answers using relevant information even when your question "
                    "is phrased differently from the document text.")

            # 0) If explicit images → VLM (with TrOCR OCR hint)
            if images and len(images) > 0:
                ocr_hints = []
                if USE_TROCR and _TROCR_AVAILABLE:
                    try:
                        for p in images:
                            with Image.open(p) as im:
                                ocr_hints.append(self._trocr_ocr_image(im) or "")
                    except Exception:
                        pass
                hint_text = "\n".join([h for h in ocr_hints if h])
                vlm_prompt = (
                    "You are a strict document QA model. Use ONLY the provided image(s). "
                    "If the answer is not clearly visible, reply exactly: \"I don't know\".\n\n"
                    f"OCR_HINT (may be partial):\n{hint_text[:4000]}\n\n"
                    f"Question: {question}"
                )
                vlm_answer = self.ask_vlm(vlm_prompt, images)
                if vlm_answer.strip():
                    ans = vlm_answer.strip()
                    if ans != "I don't know":
                        ans += note
                    return ans

            # 1) Summary path (unchanged)
            if qtype == "summary":
                return self.summarize_document(question, chat_id, document_id, combine_docs)

            # 2) TEXT RAG first (GPU text LLM) — primary for text-based docs
            context, used_docs = self.get_context(
                question=question,
                chat_id=chat_id,
                document_id=document_id,
                combine_docs=combine_docs,
                k=DEFAULT_K,
            )

            # 3) If context weak OR scanned PDF (low text) → VLM fallback
            need_auto_vlm = (len(context.strip()) < self.min_context_chars_for_text)
            if not need_auto_vlm:
                target_doc = document_id or (combine_docs[0] if combine_docs else None)
                src = self._load_source_meta(chat_id, target_doc) if target_doc else None
                if src and str(src).lower().endswith(".pdf"):
                    chars = self._probe_pdf_text_chars(src, probe_pages=3)
                    need_auto_vlm = chars < self.pdf_low_text_chars_probe

            if need_auto_vlm and AUTO_VLM_ON_LOW_TEXT:
                vlm_auto = self._auto_vlm_from_pdf(
                    question, chat_id, document_id or (combine_docs[0] if combine_docs else None)
                )
                if vlm_auto:
                    return vlm_auto

                if not context:
                    return "Not found in the provided document."

                # >>> DETERMINISTIC COUNT/LIST LOGIC (banks) <<<
                intent = self._parse_intent(question)
                if intent["wants_count"] or intent["startswith"] or intent["location"]:
                    lines = self._candidate_lines_from_context(context)
                    banks = self._filter_banks(lines, intent["startswith"], intent["location"])

                    if intent["wants_count"]:
                        if intent["location"]:
                            loc = intent["location"]
                            return f"{len(banks)} banks in {loc}" if banks else "I don't know"
                        return f"{len(banks)} banks found" if banks else "I don't know"

                    if intent["startswith"]:
                        if not banks:
                            return "I don't know"
                        show = banks[:10]
                        more = "" if len(banks) <= 10 else f"\n\n(+{len(banks)-10} more)"
                        return "- " + "\n- ".join(show) + more

            # 4) Grounded text answer via GPU text LLM
            # 4) Grounded text answer via GPU text LLM (strict, context-only)
            prompt = (
                "You are a STRICT retrieval QA model.\n"
                "RULES:\n"
                "1) Use ONLY the provided CONTEXT. Do NOT use outside knowledge.\n"
                "2) If the answer is not FULLY supported by the CONTEXT, reply exactly: \"I don't know\".\n"
                "3) Prefer copying exact numbers/names/phrases from CONTEXT.\n"
                "4) If multiple candidates exist, list them all; do not invent.\n"
                "5) Be concise.\n\n"
                "QUESTION:\n"
                f"{question}\n\n"
                "=== CONTEXT START ===\n"
                f"{context}\n"
                "=== CONTEXT END ===\n\n"
                "FINAL ANSWER ONLY (no preface/explanation):"
            )

            payload = {
                "model": self.model_name,            # qwen2.5vl-gpu:latest for both text & vision (as you set)
                "prompt": prompt,
                "stream": False,
                "keep_alive": "10m",
                "options": {
                    # very low creativity to avoid guesswork
                    "temperature": 0.0,
                    "top_p": 0.8,
                    "repeat_penalty": 1.05,
                    # keep these generous if your chunks are big
                    "num_ctx": 4096,
                    "num_predict": 400,
                    # optional: stop at double newline if your model tends to ramble
                    # "stop": ["\n\n"]
                },
            }
            r = requests.post(f"{self.ollama_url}/api/generate", json=payload, timeout=self.request_timeout)
            r.raise_for_status()
            answer = (r.json().get("response") or "").strip()
            if not answer:
                return "I don't know"

            if self.verify_pass:
                answer = self._verify_answer(question, context, answer).strip()

            if answer and answer != "I don't know":
                answer += note
            return answer


        except requests.RequestException as e:
            print(f"❌ Error calling Ollama API: {e}")
            raise
        except Exception as e:
            print(f"❌ Error in answer_question: {e}")
            raise
